"""
parser.py
Extracts plain text from uploaded resume files (PDF, DOCX, TXT)
and pulls out lightweight metadata (name guess, email, phone, skills).
"""

import os
import re

import pdfplumber
import docx


def extract_text(filepath: str) -> str:
    """Extract raw text from a resume file based on its extension."""
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".pdf":
        return _extract_pdf(filepath)
    elif ext == ".docx":
        return _extract_docx(filepath)
    elif ext == ".txt":
        return _extract_txt(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(filepath: str) -> str:
    text_chunks = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_chunks.append(page_text)
    return "\n".join(text_chunks)


def _extract_docx(filepath: str) -> str:
    document = docx.Document(filepath)
    paragraphs = [p.text for p in document.paragraphs]
    # Also grab table content, since resumes sometimes use tables for layout
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def _extract_txt(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"(\+?\d[\d\-\s().]{8,}\d)")


def guess_name(text: str, fallback: str) -> str:
    """Best-effort guess at the candidate's name: first non-empty line
    that looks like a name (no @ sign, no digits, short length)."""
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if "@" in line or any(ch.isdigit() for ch in line):
            continue
        word_count = len(line.split())
        if 1 <= word_count <= 4 and len(line) <= 50:
            return line
    return fallback


def extract_email(text: str) -> str:
    match = EMAIL_RE.search(text)
    return match.group(0) if match else "—"


def extract_phone(text: str) -> str:
    match = PHONE_RE.search(text)
    return match.group(0).strip() if match else "—"


def clean_filename_as_fallback(filename: str) -> str:
    name = os.path.splitext(filename)[0]
    name = re.sub(r"[_\-]+", " ", name)
    return name.strip().title()
